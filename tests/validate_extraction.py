"""Validate F4 text extraction and OCR adapter behavior."""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from intake import (  # noqa: E402
    DoclingBlock,
    DoclingConversion,
    DoclingTable,
    OCR_COMPLETED,
    OCR_FAILED,
    OCR_NOT_REQUIRED,
    OCR_PENDING,
    extract_document,
    extract_text,
)


class FakeOcrEngine:
    def __init__(self, text: str, *, fail: bool = False) -> None:
        self.text = text
        self.fail = fail
        self.calls: list[Path] = []

    def recognize_image(
        self,
        image_path: Path,
        *,
        languages: tuple[str, ...] | None = None,
    ) -> str:
        del languages
        self.calls.append(image_path)
        if self.fail:
            from intake import OcrRuntimeError

            raise OcrRuntimeError("fake OCR failure")
        return self.text


class FakeDocumentUnderstanding:
    def convert(self, source_path: Path) -> DoclingConversion:
        return DoclingConversion(
            text=f"Docling normalized {source_path.stem}",
            blocks=(DoclingBlock("paragraph", "structured paragraph", 1, None, "#/texts/0"),),
            tables=(DoclingTable(1, (("column",), ("value",)), "#/tables/0"),),
            page_count=1,
            extractor_version="2.41.0-test",
            model_version="fixture-model",
        )


def main() -> None:
    with tempfile.TemporaryDirectory() as temporary_dir:
        workspace = Path(temporary_dir)

        pdf_path = workspace / "hearing-notice.pdf"
        _write_pdf(pdf_path, "Hearing notice for Nairobi matter 001")
        pdf_result = extract_text(pdf_path)
        assert pdf_result.detected_file_type == "pdf"
        assert pdf_result.page_count == 1
        assert pdf_result.ocr_status == OCR_NOT_REQUIRED
        assert "Nairobi matter 001" in pdf_result.text

        docx_path = workspace / "draft-pleading.docx"
        _write_docx(docx_path, "Draft pleading for commercial division")
        docx_result = extract_text(docx_path)
        assert docx_result.detected_file_type == "docx"
        assert docx_result.ocr_status == OCR_NOT_REQUIRED
        assert "commercial division" in docx_result.text

        structured_result = extract_document(
            pdf_path,
            document_understanding=FakeDocumentUnderstanding(),
        )
        assert structured_result.text == "Docling normalized hearing-notice"
        assert structured_result.blocks[0].block_type == "paragraph"
        assert structured_result.tables[0].rows == (("column",), ("value",))
        assert structured_result.model_version == "fixture-model"

        image_path = workspace / "scan.png"
        image_path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"scanned image bytes")
        image_result = extract_text(image_path)
        assert image_result.detected_file_type == "png"
        assert image_result.text == ""
        assert image_result.page_count == 1
        assert image_result.ocr_status == OCR_PENDING
        assert "ocr_adapter_pending" in image_result.warnings

        fake_ocr = FakeOcrEngine("Scanned affidavit text from OCR")
        image_ocr_result = extract_text(image_path, ocr_engine=fake_ocr)
        assert image_ocr_result.ocr_status == OCR_COMPLETED
        assert image_ocr_result.text == "Scanned affidavit text from OCR"
        assert "ocr_text_extracted" in image_ocr_result.warnings
        assert fake_ocr.calls == [image_path]

        failing_ocr_result = extract_text(image_path, ocr_engine=FakeOcrEngine("", fail=True))
        assert failing_ocr_result.ocr_status == OCR_FAILED
        assert "ocr_failed" in failing_ocr_result.warnings

        scanned_pdf_path = workspace / "scanned-annexure.pdf"
        _write_image_only_pdf(scanned_pdf_path)
        scanned_pending = extract_text(scanned_pdf_path)
        assert scanned_pending.ocr_status == OCR_PENDING
        assert "empty_extracted_text" in scanned_pending.warnings

        scanned_ocr = FakeOcrEngine("OCR text from image only PDF annexure")
        scanned_result = extract_text(scanned_pdf_path, ocr_engine=scanned_ocr)
        assert scanned_result.detected_file_type == "pdf"
        assert scanned_result.page_count == 1
        assert scanned_result.ocr_status == OCR_COMPLETED
        assert "image only PDF annexure" in scanned_result.text
        assert "ocr_text_extracted" in scanned_result.warnings

        unsupported_path = workspace / "notes.txt"
        unsupported_path.write_text("not supported for extraction", encoding="utf-8")
        unsupported_result = extract_text(unsupported_path)
        assert unsupported_result.detected_file_type == "unsupported"
        assert unsupported_result.page_count == 0
        assert "unsupported_for_extraction" in unsupported_result.warnings

    print("EXTRACTION VALIDATION PASS")


def _write_pdf(path: Path, text: str) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    document.save(path)
    document.close()


def _write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def _write_image_only_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 320, 120), False)
    pixmap.clear_with(230)
    page.insert_image(fitz.Rect(72, 120, 523, 320), pixmap=pixmap)
    document.save(path)
    document.close()


if __name__ == "__main__":
    main()
