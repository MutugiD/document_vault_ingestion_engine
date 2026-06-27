"""Validate the private-document manual ingest smoke runner."""

from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parents[1]


def main() -> None:
    with tempfile.TemporaryDirectory() as temporary_dir:
        workspace = Path(temporary_dir)
        input_root = workspace / "private-input"
        scratch = workspace / "scratch"
        input_root.mkdir()

        sensitive_terms = (
            "Mavuno Hardware",
            "LR No. 209/152",
            "Kirinyaga Road",
            "PV-77",
            "Counterclaim Draft",
            "private-text-layer.pdf",
            "private-pleading.docx",
            "private-scanned.pdf",
            "legacy-private.doc",
        )
        _write_pdf(
            input_root / "private-text-layer.pdf",
            (
                "Mavuno Hardware records LR No. 209/152 on Kirinyaga Road with "
                "urgent rent arrears evidence and a preservation issue."
            ),
        )
        _write_docx(
            input_root / "private-pleading.docx",
            (
                "Counterclaim Draft pleads procedural unfairness, rectification, "
                "and a triable issue on allotment."
            ),
        )
        _write_scanned_pdf(input_root / "private-scanned.pdf")
        (input_root / "legacy-private.doc").write_bytes(
            b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1legacy-private"
        )

        command = [
            sys.executable,
            str(ROOT / "scripts" / "manual_ingest_smoke.py"),
            "--input",
            str(input_root),
            "--workspace",
            str(scratch),
        ]
        result = subprocess.run(command, check=False, capture_output=True, text=True)
        assert result.returncode == 0, result.stdout + result.stderr
        assert "MANUAL INGEST SMOKE PASS" in result.stdout
        assert '"duplicate_detected": true' in result.stdout
        assert '"legacy_doc_rejected": true' in result.stdout
        assert '"scanned_pdf_seen": true' in result.stdout
        for term in sensitive_terms:
            assert term not in result.stdout
            assert term not in result.stderr

    print("MANUAL INGEST SMOKE VALIDATION PASS")


def _write_pdf(path: Path, text: str) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_textbox(fitz.Rect(72, 72, 523, 760), text, fontsize=11)
    document.save(path)
    document.close()


def _write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_heading("Counterclaim Draft", level=1)
    document.add_paragraph(text)
    document.save(path)


def _write_scanned_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 320, 140), False)
    pixmap.clear_with(210)
    page.insert_image(fitz.Rect(72, 120, 523, 320), pixmap=pixmap)
    page.draw_rect(fitz.Rect(72, 120, 523, 320), color=(0, 0, 0), width=1)
    document.save(path)
    document.close()


if __name__ == "__main__":
    main()
