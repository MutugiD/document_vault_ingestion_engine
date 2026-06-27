"""Validate real-world style legal document intake, RAG, copy, vault, and backup flow."""

from __future__ import annotations

import base64
import sys
import tempfile
from dataclasses import dataclass
from datetime import UTC, date, datetime
from pathlib import Path

import fitz
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import padding, rsa
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from backup import (  # noqa: E402
    InMemoryGrantBackend,
    create_local_backup,
    create_upload_grant,
    restore_backup_package,
    upload_encrypted_snapshot,
)
from intake import (  # noqa: E402
    ACCEPTED_STATUS,
    DUPLICATE_STATUS,
    OCR_COMPLETED,
    OCR_NOT_REQUIRED,
    OCR_PENDING,
    REJECTED_STATUS,
    extract_text,
    import_document,
)
from licensing import (  # noqa: E402
    ACTIVE_STATUS,
    FeatureEntitlements,
    LicenseDocument,
    canonical_license_bytes,
    ensure_installation_identity,
    verify_license_document,
)
from rag import build_answer_packet, build_rag_index, chunk_text, retrieve_context  # noqa: E402
from search import (  # noqa: E402
    ARCHIVED_STATUS,
    DRAFT_STATUS,
    FILED_STATUS,
    REVIEWED_STATUS,
    add_document_version,
    create_document,
    create_matter,
    search_documents,
)
from vault import initialize_vault, open_vault  # noqa: E402


@dataclass(frozen=True)
class StoredFixture:
    title: str
    document_type: str
    source_path: Path
    document_id: str
    stored_object_id: str
    source_sha256: str
    extracted_text: str


class FakeOcrEngine:
    def __init__(self, text: str) -> None:
        self.text = text
        self.calls: list[Path] = []

    def recognize_image(
        self,
        image_path: Path,
        *,
        languages: tuple[str, ...] | None = None,
    ) -> str:
        del languages
        self.calls.append(image_path)
        return self.text


def main() -> None:
    vault_phrase = "real world rag validator recovery phrase"
    with tempfile.TemporaryDirectory() as temporary_dir:
        workspace = Path(temporary_dir)
        source_dir = workspace / "source-documents"
        source_dir.mkdir()
        vault_root = workspace / "vault"
        backup_path = workspace / "snapshots" / "real-world.wakilibak"
        restore_root = workspace / "restore"

        identity = ensure_installation_identity(workspace / "settings" / "installation.json")
        public_key_pem, license_document = _signed_license(identity.installation_id)
        license_result = verify_license_document(
            license_document,
            public_key_pem,
            identity.installation_id,
            as_of=date(2026, 6, 26),
        )
        assert license_result.status == ACTIVE_STATUS
        assert license_result.feature_enabled("document_intake")
        assert license_result.feature_enabled("cloud_backup")
        assert license_result.feature_enabled("managed_restore")
        assert license_result.feature_enabled("matter_rag")

        old_pdf = source_dir / "1978-archived-lease-and-rent-demand.pdf"
        pleading_docx = source_dir / "defence-and-counterclaim-draft.docx"
        scanned_pdf = source_dir / "scanned-annexure-without-ocr.pdf"
        legacy_doc = source_dir / "legacy-word-document.doc"

        old_pdf_text = _old_archive_pdf_text()
        pleading_text = _pleading_docx_text()
        _write_old_pdf(old_pdf, old_pdf_text)
        _write_docx(pleading_docx, pleading_text)
        _write_image_only_pdf(scanned_pdf)
        legacy_doc.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1legacy word bytes")

        scan_without_runtime = extract_text(scanned_pdf)
        assert scan_without_runtime.ocr_status == OCR_PENDING
        assert scan_without_runtime.text == ""

        vault_session = initialize_vault(vault_root, vault_phrase)
        matter = create_matter(
            vault_root,
            internal_reference="RW-RAG-001",
            client_name="Mavuno Hardware Ltd",
            parties="Mavuno Hardware Ltd v County Lands Registrar",
            court="Environment and Land Court",
            station="Nairobi",
            case_number="ELC RW 1978 of 2026",
            practice_area="Land and Commercial",
            responsible_advocate="M. Mutugi",
        )

        stored_old_pdf = _intake_store_and_version(
            vault_root,
            vault_session,
            matter.matter_id,
            old_pdf,
            title="Archived Lease And Rent Demand",
            document_type="Archive",
            lifecycle_status=ARCHIVED_STATUS,
        )
        stored_docx = _intake_store_and_version(
            vault_root,
            vault_session,
            matter.matter_id,
            pleading_docx,
            title="Defence And Counterclaim Draft",
            document_type="Pleading",
            lifecycle_status=REVIEWED_STATUS,
        )
        stored_scan = _intake_store_and_version(
            vault_root,
            vault_session,
            matter.matter_id,
            scanned_pdf,
            title="Scanned Annexure Pending OCR",
            document_type="Annexure",
            lifecycle_status=DRAFT_STATUS,
            ocr_engine=FakeOcrEngine("Handwritten annexure note confirms payment voucher PV-77"),
        )

        duplicate_copy = source_dir / "copy-of-1978-archived-lease.pdf"
        duplicate_copy.write_bytes(old_pdf.read_bytes())
        duplicate_record = import_document(vault_root, duplicate_copy)
        assert duplicate_record.status == DUPLICATE_STATUS
        assert "duplicate_source_hash" in duplicate_record.warnings
        assert duplicate_record.quarantine_path.read_bytes() == duplicate_copy.read_bytes()
        assert duplicate_copy.exists(), "intake must copy source documents, not move them"

        legacy_record = import_document(vault_root, legacy_doc)
        assert legacy_record.status == REJECTED_STATUS
        assert legacy_record.detected_file_type == "unsupported"

        _assert_vault_copy_and_encryption(vault_session, stored_old_pdf)
        _assert_vault_copy_and_encryption(vault_session, stored_docx)
        _assert_vault_copy_and_encryption(vault_session, stored_scan)

        expected_chunks = (
            len(chunk_text(stored_old_pdf.extracted_text))
            + len(chunk_text(stored_docx.extracted_text))
            + len(chunk_text(stored_scan.extracted_text))
        )
        indexed_count = build_rag_index(vault_root, matter_id=matter.matter_id)
        assert indexed_count == expected_chunks
        assert indexed_count >= 3

        search_hits = search_documents(
            vault_root,
            "Kirinyaga Road allotment",
            matter_id=matter.matter_id,
        )
        assert search_hits
        assert search_hits[0].matter_id == matter.matter_id

        for question, expected_phrase in _many_grounded_questions():
            grounded_query = f"{question} {expected_phrase}"
            packet = build_answer_packet(
                vault_root,
                grounded_query,
                matter_id=matter.matter_id,
                top_k=4,
            )
            assert packet.citations, question
            assert expected_phrase.lower() in packet.grounded_context.lower(), question
            assert "Use only the cited local context" in packet.safety_notice

        scan_packet = build_answer_packet(
            vault_root,
            "Which payment voucher is handwritten on the scanned annexure PV-77?",
            matter_id=matter.matter_id,
            top_k=4,
        )
        assert scan_packet.citations
        assert "PV-77" in scan_packet.grounded_context

        scan_results = retrieve_context(
            vault_root,
            "What handwritten note appears on the scanned annexure PV-77?",
            matter_id=matter.matter_id,
        )
        assert any(result.chunk.document_id == stored_scan.document_id for result in scan_results)

        backup_package = create_local_backup(
            vault_root,
            backup_path,
            recovery_key=vault_phrase,
            installation_id=identity.installation_id,
        )
        package_bytes = backup_path.read_bytes()
        assert b"Mavuno Hardware" not in package_bytes
        assert b"Kirinyaga Road" not in package_bytes
        assert b"triable issue" not in package_bytes

        backend = InMemoryGrantBackend()
        upload = upload_encrypted_snapshot(
            create_upload_grant(
                "aws",
                identity.installation_id,
                backup_package.manifest.snapshot_id,
            ),
            backup_path,
            backend=backend,
        )
        assert set(upload.metadata) == {
            "installation_id",
            "snapshot_id",
            "package_size_bytes",
            "package_hash",
            "created_at",
            "app_version",
            "upload_status",
        }
        assert "client_name" not in upload.metadata
        assert "case_number" not in upload.metadata

        restore_report = restore_backup_package(
            backup_path,
            restore_root,
            recovery_key=vault_phrase,
        )
        restored_session = open_vault(restore_report.restored_path, vault_phrase)
        assert restored_session.read_object(stored_old_pdf.stored_object_id) == old_pdf.read_bytes()
        assert (
            restored_session.read_object(stored_docx.stored_object_id) == pleading_docx.read_bytes()
        )
        assert (
            restored_session.read_object(stored_scan.stored_object_id) == scanned_pdf.read_bytes()
        )

    print("REAL WORLD RAG E2E VALIDATION PASS")


def _intake_store_and_version(
    vault_root: Path,
    vault_session,
    matter_id: str,
    source_path: Path,
    *,
    title: str,
    document_type: str,
    lifecycle_status: str,
    ocr_engine=None,
) -> StoredFixture:
    record = import_document(vault_root, source_path)
    assert record.status == ACCEPTED_STATUS
    assert record.quarantine_path.exists()
    assert record.quarantine_path != source_path
    assert record.quarantine_path.read_bytes() == source_path.read_bytes()
    assert source_path.exists(), "intake must not move or delete source documents"

    extraction = extract_text(record.quarantine_path, ocr_engine=ocr_engine)
    if source_path.name.startswith("scanned-") and ocr_engine is None:
        assert extraction.ocr_status == OCR_PENDING
        assert extraction.text == ""
        assert "empty_extracted_text" in extraction.warnings
    else:
        assert extraction.ocr_status in {OCR_NOT_REQUIRED, OCR_COMPLETED}
        assert extraction.text

    stored_object = vault_session.write_object(
        record.quarantine_path.read_bytes(),
        original_name=source_path.name,
        content_type=_content_type(record.detected_file_type),
        actor="real-world-validator",
    )
    document = create_document(
        vault_root,
        matter_id=matter_id,
        title=title,
        document_type=document_type,
        lifecycle_status=FILED_STATUS if lifecycle_status == FILED_STATUS else DRAFT_STATUS,
    )
    add_document_version(
        vault_root,
        document_id=document.document_id,
        object_id=stored_object.object_id,
        source_sha256=stored_object.sha256,
        extracted_text=extraction.text,
        lifecycle_status=lifecycle_status,
    )
    return StoredFixture(
        title=title,
        document_type=document_type,
        source_path=source_path,
        document_id=document.document_id,
        stored_object_id=stored_object.object_id,
        source_sha256=stored_object.sha256,
        extracted_text=extraction.text,
    )


def _assert_vault_copy_and_encryption(vault_session, fixture: StoredFixture) -> None:
    stored = vault_session.get_object(fixture.stored_object_id)
    source_bytes = fixture.source_path.read_bytes()
    assert stored.sha256 == fixture.source_sha256
    assert vault_session.read_object(fixture.stored_object_id) == source_bytes
    assert stored.object_path.exists()
    assert stored.object_path.read_bytes() != source_bytes
    assert fixture.source_path.name not in stored.object_path.name


def _many_grounded_questions() -> tuple[tuple[str, str], ...]:
    old_pdf_questions = (
        ("Which parcel number appears in the old archive?", "LR No. 209/152"),
        ("What road is mentioned in the archived lease?", "Kirinyaga Road"),
        ("What rent arrears amount is demanded?", "KES 84,500"),
        ("Which allotment letter date is cited?", "12 March 1978"),
        ("Which receipts support the historical rent payments?", "receipt numbers 4412 and 4413"),
        ("What covenant is relied on for quiet possession?", "quiet possession"),
        ("Which notice period appears in the rent demand?", "fourteen day notice"),
        ("What registry problem is described?", "missing green card entry"),
        ("Which survey plan is referenced?", "Survey Plan 91/14"),
        ("Who signed the archive minute?", "Deputy Registrar M. Wekesa"),
        ("What business operated from the premises?", "hardware storage business"),
        ("What old file number is referenced?", "Archive File 1978/44"),
        ("What valuation phrase appears in the archive?", "unimproved site value"),
        ("Which party requested reconstruction?", "Mavuno Hardware Ltd"),
        ("What registry station is named?", "Nairobi registry"),
        ("Which rent year is disputed?", "1979"),
        ("What annexure label is used for the lease?", "Annexure ML-1"),
        ("What was the lease term?", "thirty three years"),
        ("Which municipal reference appears?", "City Council rent register"),
        ("What does the archive say about transfer consent?", "consent was endorsed"),
    )
    pleading_questions = (
        ("What is the main defence to eviction?", "procedural unfairness"),
        ("Which issue is said to be triable?", "triable issue on allotment"),
        ("What counterclaim relief is requested?", "rectification of the registry"),
        ("Which doctrine is mentioned on limitation?", "continuing breach"),
        ("What affidavit witness is proposed?", "Joseph Nderitu"),
        ("What order is sought on status quo?", "preservation of the premises"),
        ("Which filing deadline is tracked?", "21 days"),
        ("What evidence supports occupation?", "stock ledgers and rent receipts"),
        (
            "What pleading section discusses jurisdiction?",
            "Environment and Land Court jurisdiction",
        ),
        ("What hearing preparation item is listed?", "bundle index"),
        ("What risk is stated if eviction proceeds?", "irreparable disruption"),
        ("What draft prayer seeks costs?", "costs of the counterclaim"),
        ("What phrase describes the county records?", "incomplete registry trail"),
        ("What is requested before cross-examination?", "discovery of the rent register"),
        ("What annexure is tied to photographs?", "Annexure MN-4"),
        ("Which witness role is described?", "records clerk"),
        ("What settlement option is preserved?", "without prejudice negotiation"),
        ("What legal threshold is mentioned?", "prima facie case"),
        ("What does the pleading say about service?", "service was disputed"),
        ("What document should be filed after review?", "amended defence"),
    )
    return old_pdf_questions + pleading_questions


def _old_archive_pdf_text() -> str:
    return (
        "Archive File 1978/44 concerns Mavuno Hardware Ltd and LR No. 209/152 on "
        "Kirinyaga Road. The allotment letter dated 12 March 1978 granted a thirty "
        "three years lease for a hardware storage business. Annexure ML-1 is the "
        "lease. The City Council rent register records receipt numbers 4412 and "
        "4413, but the 1979 rent year is disputed. A rent demand claims KES 84,500 "
        "and gives a fourteen day notice. The archive mentions quiet possession, "
        "unimproved site value, Survey Plan 91/14, Nairobi registry, and a missing "
        "green card entry. Mavuno Hardware Ltd requested reconstruction because "
        "transfer consent was endorsed. The archive minute was signed by Deputy "
        "Registrar M. Wekesa."
    )


def _pleading_docx_text() -> str:
    return (
        "The defence and counterclaim pleads procedural unfairness and a triable "
        "issue on allotment. It seeks rectification of the registry, preservation "
        "of the premises, and costs of the counterclaim. The limitation answer is "
        "continuing breach. Evidence includes stock ledgers and rent receipts. "
        "The pleading relies on Environment and Land Court jurisdiction, a prima "
        "facie case, and irreparable disruption if eviction proceeds. Joseph "
        "Nderitu is proposed as a witness and records clerk. Service was disputed. "
        "The county records show an incomplete registry trail. Discovery of the "
        "rent register is requested before cross-examination. Annexure MN-4 covers "
        "photographs. The filing deadline is 21 days. The hearing preparation item "
        "is a bundle index. A without prejudice negotiation remains open. After "
        "review, the amended defence should be filed."
    )


def _write_old_pdf(path: Path, text: str) -> None:
    document = fitz.open()
    document.set_metadata(
        {
            "title": "Archived Lease And Rent Demand",
            "author": "Registry archive",
            "creationDate": "D:19780312090000+03'00'",
        }
    )
    first, second = text[: len(text) // 2], text[len(text) // 2 :]
    for page_text in (first, second):
        page = document.new_page(width=595, height=842)
        page.insert_textbox(
            fitz.Rect(72, 72, 523, 760),
            page_text,
            fontsize=11,
            fontname="helv",
        )
    document.save(path)
    document.close()


def _write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_heading("Defence And Counterclaim Draft", level=1)
    for sentence in text.split(". "):
        document.add_paragraph(sentence.strip().rstrip(".") + ".")
    document.save(path)


def _write_image_only_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 360, 160), False)
    pixmap.clear_with(230)
    page.insert_image(fitz.Rect(72, 120, 523, 320), pixmap=pixmap)
    page.draw_rect(fitz.Rect(72, 120, 523, 320), color=(0, 0, 0), width=1)
    document.save(path)
    document.close()


def _content_type(detected_file_type: str) -> str:
    return {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "png": "image/png",
        "jpeg": "image/jpeg",
        "tiff": "image/tiff",
    }[detected_file_type]


def _signed_license(installation_id: str) -> tuple[bytes, LicenseDocument]:
    private_key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    public_key_pem = private_key.public_key().public_bytes(
        encoding=serialization.Encoding.PEM,
        format=serialization.PublicFormat.SubjectPublicKeyInfo,
    )
    unsigned_document = LicenseDocument(
        installation_id=installation_id,
        license_id="LIC-REAL-WORLD-RAG",
        firm_display_name="Example Advocates LLP",
        plan="firm",
        features=FeatureEntitlements(
            document_intake=True,
            cloud_backup=True,
            managed_restore=True,
            matter_rag=True,
        ),
        expiry=date(2099, 12, 31),
        issued_at=datetime(2026, 6, 26, 11, 0, tzinfo=UTC),
        signature="",
    )
    signature = private_key.sign(
        canonical_license_bytes(unsigned_document),
        padding.PSS(
            mgf=padding.MGF1(hashes.SHA256()),
            salt_length=padding.PSS.MAX_LENGTH,
        ),
        hashes.SHA256(),
    )
    return public_key_pem, LicenseDocument(
        installation_id=unsigned_document.installation_id,
        license_id=unsigned_document.license_id,
        firm_display_name=unsigned_document.firm_display_name,
        plan=unsigned_document.plan,
        features=unsigned_document.features,
        expiry=unsigned_document.expiry,
        issued_at=unsigned_document.issued_at,
        signature=base64.b64encode(signature).decode("ascii"),
    )


if __name__ == "__main__":
    main()
