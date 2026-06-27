"""Validate public Kenyan legal-document E2E runner with deterministic fixtures."""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parents[1]


def main() -> None:
    with tempfile.TemporaryDirectory() as temporary_dir:
        workspace = Path(temporary_dir)
        input_root = workspace / "public-docs"
        input_root.mkdir()
        _write_pdf(
            input_root / "supreme-court-rules.pdf",
            "The Supreme Court rules provide forms, notices, petitions, and procedure "
            "for filing applications in the Supreme Court of Kenya.",
        )
        _write_pdf(
            input_root / "land-commission-stay.pdf",
            "The National Land Commission stay application concerns land, registry "
            "procedure, court orders, and a public law dispute.",
        )
        _write_pdf(
            input_root / "court-of-appeal-registry.pdf",
            "The Court of Appeal registry manual discusses appeals, notices, records, "
            "registry workflow, and court filing procedure.",
        )
        _write_docx(
            input_root / "self-represented-litigant-guide.docx",
            "The self represented litigant guide explains affidavits, pleadings, forms, "
            "petitioners, notices, and electronic case management steps.",
        )
        _write_image_only_pdf(input_root / "scanned-practice-direction.pdf")
        (input_root / "scanned-practice-direction.pdf.ocr.txt").write_text(
            "Scanned public court practice direction mentions automation, registry intake, "
            "virtual court workflow, and case management.",
            encoding="utf-8",
        )
        (input_root / "duplicate-supreme-court-rules.pdf").write_bytes(
            (input_root / "supreme-court-rules.pdf").read_bytes()
        )
        (input_root / "legacy-public-notice.doc").write_bytes(
            b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1legacy public notice bytes"
        )

        report_path = workspace / "report.json"
        command = [
            sys.executable,
            str(ROOT / "scripts" / "public_kenyan_e2e.py"),
            "--input",
            str(input_root),
            "--workspace",
            str(workspace / "scratch"),
            "--report",
            str(report_path),
        ]
        result = subprocess.run(command, check=False, capture_output=True, text=True)
        assert result.returncode == 0, result.stdout + result.stderr
        assert "PUBLIC KENYAN E2E PASS" in result.stdout
        assert '"restore_verified": true' in result.stdout
        assert '"confidence":' in result.stdout
        assert '"answer_status": "grounded_context_available"' in result.stdout
        assert "provide forms, notices, petitions" not in result.stdout
        assert "registry workflow, and court filing procedure" not in result.stdout
        assert report_path.exists()
        report = json.loads(report_path.read_text(encoding="utf-8"))
        assert report["indexed_documents"] >= 5
        assert report["accepted_by_type"]["pdf"] >= 4
        assert report["accepted_by_type"]["docx"] == 1
        assert report["duplicate_count"] == 1
        assert report["rejected_unsupported_count"] == 1
        assert report["scanned_ocr_completed_count"] == 1
        assert report["intake_copy_verified_count"] >= 7
        assert len(report["answers"]) == 10

        app_command = [
            sys.executable,
            str(ROOT / "main.py"),
            "--public-kenya-e2e",
            str(input_root),
        ]
        app_result = subprocess.run(app_command, check=False, capture_output=True, text=True)
        assert app_result.returncode == 0, app_result.stdout + app_result.stderr
        assert '"answers":' in app_result.stdout
        assert '"citations":' in app_result.stdout
        assert '"confidence":' in app_result.stdout
        assert "provide forms, notices, petitions" not in app_result.stdout

    print("PUBLIC KENYAN E2E VALIDATION PASS")


def _write_pdf(path: Path, text: str) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_textbox(fitz.Rect(72, 72, 523, 760), text, fontsize=11)
    document.save(path)
    document.close()


def _write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_heading("Self Represented Litigant Guide", level=1)
    for sentence in text.split(". "):
        document.add_paragraph(sentence.strip().rstrip(".") + ".")
    document.save(path)


def _write_image_only_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 320, 140), False)
    pixmap.clear_with(220)
    page.insert_image(fitz.Rect(72, 120, 523, 320), pixmap=pixmap)
    page.draw_rect(fitz.Rect(72, 120, 523, 320), color=(0, 0, 0), width=1)
    document.save(path)
    document.close()


if __name__ == "__main__":
    main()
