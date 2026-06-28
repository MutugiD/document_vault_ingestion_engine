"""Create a reproducible public/sample document corpus for manual Windows E2E."""

from __future__ import annotations

import argparse
from pathlib import Path

import fitz
from docx import Document


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Create manual E2E public/sample corpus.")
    parser.add_argument(
        "--output",
        required=True,
        type=Path,
        help="Folder where generated public/sample documents should be written.",
    )
    args = parser.parse_args(argv)
    sources = create_manual_e2e_corpus(args.output)
    print("MANUAL E2E CORPUS CREATED")
    print(f"Document count: {len(sources)}")
    print(f"Output: {args.output}")
    return 0


def create_manual_e2e_corpus(output: Path) -> list[Path]:
    output.mkdir(parents=True, exist_ok=True)
    sources: list[Path] = []
    for index in range(1, 46):
        path = output / f"public-kenyan-sample-{index:02d}.pdf"
        _write_pdf(
            path,
            (
                f"Manual corpus topic {index} discusses Kenyan court filing, registry "
                f"procedure, affidavits, pleadings, citations, and evidence item {index}."
            ),
        )
        sources.append(path)
    for index in range(46, 53):
        path = output / f"public-kenyan-drafting-sample-{index:02d}.docx"
        _write_docx(
            path,
            (
                f"Manual corpus topic {index} covers legal drafting, witness statements, "
                f"case management, hearing notices, and court registry evidence {index}."
            ),
        )
        sources.append(path)

    scanned = output / "scanned-public-practice-direction.pdf"
    _write_image_only_pdf(scanned)
    scanned.with_suffix(".pdf.ocr.txt").write_text(
        "Manual corpus scanned practice direction covers virtual court workflow and OCR evidence.",
        encoding="utf-8",
    )
    sources.append(scanned)

    duplicate = output / "duplicate-public-kenyan-sample-01.pdf"
    duplicate.write_bytes(sources[0].read_bytes())
    sources.append(duplicate)

    legacy = output / "legacy-public-form.doc"
    legacy.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1legacy public form bytes")
    sources.append(legacy)

    empty = output / "empty-public-filing.pdf"
    empty.write_bytes(b"")
    sources.append(empty)

    corrupt = output / "corrupt-public-filing.pdf"
    corrupt.write_bytes(b"%PDF-1.7 corrupt public filing bytes")
    sources.append(corrupt)
    return sources


def _write_pdf(path: Path, text: str) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_textbox(fitz.Rect(72, 72, 523, 760), text, fontsize=11)
    document.save(path)
    document.close()


def _write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_heading("Manual Public Legal Sample", level=1)
    document.add_paragraph(text)
    document.save(path)


def _write_image_only_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 320, 140), False)
    pixmap.clear_with(220)
    page.insert_image(fitz.Rect(72, 120, 523, 320), pixmap=pixmap)
    page.draw_rect(fitz.Rect(72, 120, 523, 320), color=(0, 0, 0), width=1)
    document.save(path)
    document.close()


if __name__ == "__main__":
    raise SystemExit(main())
